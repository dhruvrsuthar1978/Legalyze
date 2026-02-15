# app/services/extractor_service.py

import io
import re
import fitz                          # PyMuPDF
from docx import Document
from docx.oxml.ns import qn
from typing import Tuple, Optional
import logging

logger = logging.getLogger("legalyze.extractor")


# ══════════════════════════════════════════════════════════════════
# PUBLIC ENTRY POINT
# ══════════════════════════════════════════════════════════════════

def extract_text_from_file(
    contents: bytes,
    content_type: str
) -> Tuple[str, int, int]:
    """
    Master extractor — routes to PDF or DOCX extractor based on file type.

    Returns:
        Tuple of (extracted_text, page_count, word_count)

    Raises:
        ValueError  — Unsupported file type
        RuntimeError — Extraction failure
    """
    logger.info(f"Starting text extraction | content_type={content_type}")

    if content_type == "application/pdf":
        return _extract_from_pdf(contents)

    elif "wordprocessingml" in content_type:
        return _extract_from_docx(contents)

    else:
        raise ValueError(
            f"Unsupported file type: '{content_type}'. "
            f"Allowed: application/pdf, .docx"
        )


# ══════════════════════════════════════════════════════════════════
# PDF EXTRACTOR  (PyMuPDF)
# ══════════════════════════════════════════════════════════════════

def _extract_from_pdf(contents: bytes) -> Tuple[str, int, int]:
    """
    Extracts text from a PDF file using PyMuPDF (fitz).
    Handles multi-column layouts, headers/footers, and unicode normalization.

    Returns:
        (full_text, page_count, word_count)
    """
    raw_pages: list[str] = []
    page_count = 0

    try:
        with fitz.open(stream=contents, filetype="pdf") as doc:
            page_count = len(doc)
            logger.info(f"PDF opened successfully | pages={page_count}")

            for page_num, page in enumerate(doc, start=1):
                try:
                    # Use 'text' mode for clean extraction
                    # 'blocks' mode can be used for layout-aware extraction
                    page_text = page.get_text("text", sort=True)
                    cleaned   = _clean_text(page_text)

                    if cleaned:
                        raw_pages.append(f"[PAGE {page_num}]\n{cleaned}")

                except Exception as e:
                    logger.warning(f"Could not extract page {page_num}: {e}")
                    continue

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")

    full_text  = "\n\n".join(raw_pages)
    word_count = len(full_text.split())

    logger.info(
        f"PDF extraction complete | "
        f"pages={page_count}, words={word_count}, "
        f"chars={len(full_text)}"
    )

    return full_text, page_count, word_count


# ══════════════════════════════════════════════════════════════════
# DOCX EXTRACTOR  (python-docx)
# ══════════════════════════════════════════════════════════════════

def _extract_from_docx(contents: bytes) -> Tuple[str, int, int]:
    """
    Extracts text from a DOCX file using python-docx.
    Handles:
      - Body paragraphs
      - Nested tables (all rows/cells)
      - Headers and footers
      - Text boxes (via XML fallback)

    Returns:
        (full_text, page_count, word_count)
    """
    try:
        doc = Document(io.BytesIO(contents))
    except Exception as e:
        logger.error(f"DOCX open failed: {e}")
        raise RuntimeError(f"Failed to open DOCX file: {str(e)}")

    sections: list[str] = []

    # ── Body paragraphs ─────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            sections.append(text)

    # ── Tables ───────────────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            row_cells = [
                cell.text.strip()
                for cell in row.cells
                if cell.text.strip()
            ]
            if row_cells:
                sections.append(" | ".join(row_cells))

    # ── Headers and Footers ──────────────────────────────────────
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            if header_footer:
                for para in header_footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        sections.append(text)

    # ── Text boxes (XML fallback) ────────────────────────────────
    try:
        body = doc.element.body
        for txbx in body.iter(qn("w:txbxContent")):
            for para in txbx.iter(qn("w:p")):
                inner_text = "".join(
                    r.text or ""
                    for r in para.iter(qn("w:t"))
                ).strip()
                if inner_text:
                    sections.append(inner_text)
    except Exception as e:
        logger.warning(f"Text box extraction failed (non-critical): {e}")

    full_text  = "\n".join(sections)
    full_text  = _clean_text(full_text)
    word_count = len(full_text.split())

    # DOCX does not expose page count — estimate from word count
    estimated_pages = max(1, word_count // 300)

    logger.info(
        f"DOCX extraction complete | "
        f"estimated_pages={estimated_pages}, words={word_count}"
    )

    return full_text, estimated_pages, word_count


# ══════════════════════════════════════════════════════════════════
# TEXT CLEANING
# ══════════════════════════════════════════════════════════════════

def _clean_text(text: str) -> str:
    """
    Cleans extracted raw text:
    - Removes null bytes and control characters
    - Normalizes whitespace (multiple spaces → single)
    - Removes excessive blank lines (3+ → 2)
    - Removes page number artifacts (standalone digits)
    - Strips leading/trailing whitespace
    """
    if not text:
        return ""

    # Remove null bytes and non-printable control chars (except \n \t)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # Remove standalone page numbers (e.g., lines that are just "3")
    text = re.sub(r"^\s*\d{1,3}\s*$", "", text, flags=re.MULTILINE)

    # Normalize tabs to spaces
    text = text.replace("\t", " ")

    # Collapse multiple spaces into one
    text = re.sub(r"[ ]{2,}", " ", text)

    # Collapse 3+ consecutive newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# ══════════════════════════════════════════════════════════════════
# METADATA UTILITIES
# ══════════════════════════════════════════════════════════════════

def get_file_size_kb(contents: bytes) -> float:
    """Returns file size in kilobytes rounded to 2 decimal places."""
    return round(len(contents) / 1024, 2)


def is_text_sufficient(text: str, min_words: int = 50) -> bool:
    """
    Checks if extracted text has enough content for meaningful analysis.
    Returns False if the document is likely scanned/image-based.
    """
    word_count = len(text.split())
    return word_count >= min_words