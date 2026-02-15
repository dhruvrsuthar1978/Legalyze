# app/services/generation_service.py

import io
import os
import uuid
from datetime import datetime
from typing import List, Tuple, Dict, Optional
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    Table, TableStyle, HRFlowable, PageBreak
)
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

logger = logging.getLogger("legalyze.generation")

OUTPUT_DIR = "generated/contracts"
os.makedirs(OUTPUT_DIR, exist_ok=True)

RISK_COLOR_MAP = {
    "High":   colors.HexColor("#FF4444"),
    "Medium": colors.HexColor("#FF8800"),
    "Low":    colors.HexColor("#22BB44"),
}


# ══════════════════════════════════════════════════════════════════
# PUBLIC ENTRY POINT
# ══════════════════════════════════════════════════════════════════

def generate_contract_document(
    original_contract: dict,
    clauses: List[dict],
    accepted_clauses: List[dict],
    format: str = "pdf",
    include_summary: bool = True,
    version: int = 1
) -> Tuple[bytes, str, float]:
    """
    Generates a new contract document incorporating accepted AI suggestions.

    Args:
        original_contract : MongoDB contract document
        clauses           : All analyzed clauses
        accepted_clauses  : Clauses with suggestion_status='accepted'/'edited'
        format            : 'pdf' or 'docx'
        include_summary   : Append AI analysis summary page
        version           : Version number for filename

    Returns:
        Tuple of (file_bytes, filename, file_size_kb)
    """
    # Build replacement map: clause_id → effective suggestion text
    replacement_map = _build_replacement_map(accepted_clauses)

    # Build final clause list with replacements applied
    final_clauses = _apply_replacements(clauses, replacement_map)

    logger.info(
        f"Generating contract | "
        f"format={format}, "
        f"total_clauses={len(final_clauses)}, "
        f"replacements={len(replacement_map)}, "
        f"include_summary={include_summary}"
    )

    if format == "pdf":
        file_bytes = _generate_pdf(
            original_contract=original_contract,
            final_clauses=final_clauses,
            replacement_map=replacement_map,
            include_summary=include_summary,
            version=version
        )
    else:
        file_bytes = _generate_docx(
            original_contract=original_contract,
            final_clauses=final_clauses,
            replacement_map=replacement_map,
            include_summary=include_summary,
            version=version
        )

    filename    = _build_filename(
        original_contract.get("filename", "contract"),
        format,
        version
    )
    file_size   = round(len(file_bytes) / 1024, 2)

    logger.info(
        f"Contract generation complete | "
        f"filename={filename}, size={file_size}KB"
    )

    return file_bytes, filename, file_size


# ══════════════════════════════════════════════════════════════════
# PDF GENERATION  (ReportLab)
# ══════════════════════════════════════════════════════════════════

def _generate_pdf(
    original_contract: dict,
    final_clauses: List[dict],
    replacement_map: Dict[str, str],
    include_summary: bool,
    version: int
) -> bytes:
    buffer  = io.BytesIO()
    styles  = getSampleStyleSheet()
    story   = []

    # ── Custom Styles ──────────────────────────────────────────
    title_style = ParagraphStyle(
        "LegalTitle",
        parent=styles["Title"],
        fontSize=22,
        textColor=colors.HexColor("#1A1A2E"),
        spaceAfter=12,
        alignment=1   # Centre
    )
    heading2 = ParagraphStyle(
        "Heading2Custom",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#16213E"),
        spaceBefore=14,
        spaceAfter=6
    )
    body_style = ParagraphStyle(
        "LegalBody",
        parent=styles["BodyText"],
        fontSize=10,
        leading=16,
        textColor=colors.HexColor("#2D2D2D"),
        spaceAfter=8
    )
    ai_style = ParagraphStyle(
        "AIReplacement",
        parent=styles["BodyText"],
        fontSize=10,
        leading=16,
        textColor=colors.HexColor("#006400"),
        backColor=colors.HexColor("#F0FFF0"),
        borderColor=colors.HexColor("#22BB44"),
        borderWidth=1,
        borderPadding=6,
        spaceAfter=8
    )
    label_style = ParagraphStyle(
        "LabelStyle",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#666666"),
        spaceAfter=4
    )

    # ── Page 1: Cover ──────────────────────────────────────────
    story.append(Spacer(1, 2 * cm))
    story.append(Paragraph("LEGALYZE", label_style))
    story.append(Paragraph("AI-REVIEWED CONTRACT", title_style))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1A1A2E")))
    story.append(Spacer(1, 0.5 * cm))

    meta_data = [
        ["Original File:",    original_contract.get("filename", "N/A")],
        ["Document Version:", f"v{version}"],
        ["Category:",         original_contract.get("category", "N/A")],
        ["Generated On:",     datetime.utcnow().strftime("%B %d, %Y at %H:%M UTC")],
        ["AI Replacements:",  str(len(replacement_map))],
        ["Total Clauses:",    str(len(final_clauses))],
    ]

    meta_table = Table(meta_data, colWidths=[5 * cm, 11 * cm])
    meta_table.setStyle(TableStyle([
        ("FONTNAME",    (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME",    (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE",    (0, 0), (-1, -1), 10),
        ("TEXTCOLOR",   (0, 0), (0, -1), colors.HexColor("#333333")),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1),
         [colors.HexColor("#F9F9F9"), colors.white]),
        ("GRID",        (0, 0), (-1, -1), 0.5, colors.HexColor("#DDDDDD")),
        ("TOPPADDING",  (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 0.5 * cm))

    story.append(Paragraph(
        "⚠️  This document has been AI-reviewed. It does not constitute legal advice. "
        "Please consult a qualified legal professional before execution.",
        ParagraphStyle(
            "Disclaimer",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.HexColor("#AA6600"),
            backColor=colors.HexColor("#FFFBEA"),
            borderPadding=8
        )
    ))
    story.append(PageBreak())

    # ── Clause Pages ──────────────────────────────────────────
    story.append(Paragraph("CONTRACT CLAUSES", title_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#CCCCCC")))
    story.append(Spacer(1, 0.3 * cm))

    for i, clause in enumerate(final_clauses, start=1):
        clause_type = clause.get("clause_type", "General")
        is_replaced = clause.get("_replaced", False)
        risk_level  = clause.get("risk_level", "Low")
        risk_color  = RISK_COLOR_MAP.get(risk_level, colors.black)

        # Clause header
        story.append(Paragraph(
            f"<b>{i}. {clause_type}</b> &nbsp;"
            f'<font color="{risk_color.hexval() if hasattr(risk_color, "hexval") else "#000000"}">'
            f"[{risk_level} Risk]</font>",
            heading2
        ))

        if is_replaced:
            story.append(Paragraph(
                f"✦ AI-Improved Clause",
                label_style
            ))
            story.append(Paragraph(clause["display_text"], ai_style))
        else:
            story.append(Paragraph(clause["display_text"], body_style))

        story.append(Spacer(1, 0.2 * cm))

    # ── Optional: Summary Page ────────────────────────────────
    if include_summary:
        story.extend(
            _build_pdf_summary_page(final_clauses, replacement_map, styles)
        )

    # ── Build PDF ─────────────────────────────────────────────
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm
    )
    doc.build(story)

    return buffer.getvalue()


def _build_pdf_summary_page(
    final_clauses: List[dict],
    replacement_map: Dict[str, str],
    styles
) -> list:
    """Builds the AI Summary report page appended to the PDF."""
    elements = [PageBreak()]
    title_style = ParagraphStyle(
        "SummaryTitle",
        parent=styles["Title"],
        fontSize=18,
        textColor=colors.HexColor("#1A1A2E"),
        alignment=1
    )
    elements.append(Paragraph("AI ANALYSIS SUMMARY", title_style))
    elements.append(HRFlowable(
        width="100%", thickness=2,
        color=colors.HexColor("#1A1A2E")
    ))
    elements.append(Spacer(1, 0.4 * cm))

    high   = sum(1 for c in final_clauses if c.get("risk_level") == "High")
    medium = sum(1 for c in final_clauses if c.get("risk_level") == "Medium")
    low    = sum(1 for c in final_clauses if c.get("risk_level") == "Low")
    total  = len(final_clauses)

    summary_data = [
        ["Metric", "Value"],
        ["Total Clauses Analyzed", str(total)],
        ["High Risk Clauses",      str(high)],
        ["Medium Risk Clauses",    str(medium)],
        ["Low Risk Clauses",       str(low)],
        ["AI Replacements Applied", str(len(replacement_map))],
        ["Clauses Retained Original", str(total - len(replacement_map))],
    ]

    table = Table(summary_data, colWidths=[9 * cm, 7 * cm])
    table.setStyle(TableStyle([
        ("BACKGROUND",  (0, 0), (-1, 0), colors.HexColor("#1A1A2E")),
        ("TEXTCOLOR",   (0, 0), (-1, 0), colors.white),
        ("FONTNAME",    (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.HexColor("#F5F5F5"), colors.white]),
        ("GRID",        (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("TOPPADDING",  (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("ALIGN",       (1, 0), (1, -1), "CENTER"),
    ]))
    elements.append(table)

    return elements


# ══════════════════════════════════════════════════════════════════
# DOCX GENERATION  (python-docx)
# ══════════════════════════════════════════════════════════════════

def _generate_docx(
    original_contract: dict,
    final_clauses: List[dict],
    replacement_map: Dict[str, str],
    include_summary: bool,
    version: int
) -> bytes:
    doc = Document()

    # ── Page margins ──────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────
    title = doc.add_heading("AI-REVIEWED CONTRACT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph(
        f"Original: {original_contract.get('filename', 'N/A')}  |  "
        f"Version: v{version}  |  "
        f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}"
    ).style = "Caption"

    doc.add_paragraph(
        "⚠️ This document has been AI-reviewed and does not constitute legal advice."
    ).runs[0].font.color.rgb = RGBColor(0xAA, 0x66, 0x00)

    doc.add_page_break()

    # ── Clause sections ───────────────────────────────────────
    doc.add_heading("CONTRACT CLAUSES", level=1)

    for i, clause in enumerate(final_clauses, start=1):
        clause_type = clause.get("clause_type", "General")
        is_replaced = clause.get("_replaced", False)
        risk_level  = clause.get("risk_level", "Low")

        heading = doc.add_heading(
            f"{i}. {clause_type}  [{risk_level} Risk]",
            level=2
        )

        # Color heading by risk level
        risk_hex = {
            "High": (0xFF, 0x44, 0x44),
            "Medium": (0xFF, 0x88, 0x00),
            "Low": (0x22, 0xBB, 0x44)
        }.get(risk_level, (0x00, 0x00, 0x00))
        heading.runs[0].font.color.rgb = RGBColor(*risk_hex)

        if is_replaced:
            label = doc.add_paragraph("✦ AI-Improved Clause")
            label.runs[0].font.italic = True
            label.runs[0].font.size   = Pt(9)
            label.runs[0].font.color.rgb = RGBColor(0x00, 0x64, 0x00)

            para = doc.add_paragraph(clause["display_text"])
            for run in para.runs:
                run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)
        else:
            doc.add_paragraph(clause["display_text"])

    # ── Summary page ──────────────────────────────────────────
    if include_summary:
        doc.add_page_break()
        doc.add_heading("AI ANALYSIS SUMMARY", level=1)

        high   = sum(1 for c in final_clauses if c.get("risk_level") == "High")
        medium = sum(1 for c in final_clauses if c.get("risk_level") == "Medium")
        low    = sum(1 for c in final_clauses if c.get("risk_level") == "Low")
        total  = len(final_clauses)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Value"

        rows = [
            ("Total Clauses", str(total)),
            ("High Risk",     str(high)),
            ("Medium Risk",   str(medium)),
            ("Low Risk",      str(low)),
            ("AI Replacements", str(len(replacement_map))),
        ]
        for label, value in rows:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value

    # ── Serialize to bytes ────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ══════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════

def _build_replacement_map(accepted_clauses: List[dict]) -> Dict[str, str]:
    """
    Builds {clause_id → effective_suggestion_text} for accepted/edited clauses.
    Edited text takes priority over raw AI suggestion.
    """
    replacement_map = {}
    for clause in accepted_clauses:
        if clause.get("suggestion_status") in ["accepted", "edited"]:
            effective = (
                clause.get("edited_suggestion")
                or clause.get("suggestion")
                or clause.get("original_text")
            )
            replacement_map[clause["clause_id"]] = effective
    return replacement_map


def _apply_replacements(
    clauses: List[dict],
    replacement_map: Dict[str, str]
) -> List[dict]:
    """
    Builds a display-ready clause list.
    Replaced clauses get 'display_text' = suggestion, '_replaced' = True.
    Original clauses keep 'display_text' = original_text, '_replaced' = False.
    """
    final = []
    for clause in clauses:
        cid = clause.get("clause_id")
        if cid and cid in replacement_map:
            final.append({
                **clause,
                "display_text": replacement_map[cid],
                "_replaced": True
            })
        else:
            final.append({
                **clause,
                "display_text": clause.get("original_text", ""),
                "_replaced": False
            })
    return final


def _build_filename(
    original_filename: str,
    format: str,
    version: int
) -> str:
    base = os.path.splitext(original_filename)[0]
    base = "".join(c if c.isalnum() or c in "-_" else "_" for c in base)
    return f"{base}_v{version}_legalyze.{format}"